from __future__ import annotations

import json
import math
from datetime import UTC, datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUTS = REPO / "outputs"
TABLES = OUTPUTS / "tables"
FIGURES = OUTPUTS / "figures"
PAPER_DIR = REPO / "paper"

NAVY = "17324D"
TEAL = "1B7F79"
INK = "20262E"
MUTED = "5F6B76"
LIGHT = "EEF3F6"
GRID = "CBD5DC"
WHITE = "FFFFFF"
ACCENT = "E0A63B"

CONTENT_WIDTH_IN = 6.9
CONTENT_WIDTH_DXA = 9936
TABLE_INDENT_DXA = 100


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, *, size=None, bold=None, italic=None, color=INK, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def shade(element, fill: str) -> None:
    properties = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to the content width")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run(run, size=8, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    properties.append(size)
    run.append(properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color, before, after in [
        ("Title", 23, NAVY, 0, 5),
        ("Subtitle", 11, MUTED, 0, 8),
        ("Heading 1", 13.5, NAVY, 9, 4),
        ("Heading 2", 11, TEAL, 6, 3),
        ("Heading 3", 9.8, NAVY, 5, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.80)
    section.bottom_margin = Inches(0.80)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    table.autofit = False
    table.columns[0].width = Inches(4.9)
    table.columns[1].width = Inches(2.0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(left.add_run("Architecture × Budget Crossover"), size=8, bold=True, color=NAVY)
    set_run(right.add_run("Technical white paper"), size=8, color=MUTED)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
    first = header.paragraphs[0]
    first._element.getparent().remove(first._element)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("Page "), size=8, color=MUTED)
    add_page_field(paragraph, "PAGE")
    set_run(paragraph.add_run(" of "), size=8, color=MUTED)
    add_page_field(paragraph, "NUMPAGES")


def add_title_block(document: Document) -> None:
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    set_run(kicker.add_run("TECHNICAL WHITE PAPER"), size=8.5, bold=True, color=ACCENT)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        title.add_run(
            "Budget Thresholds for LLM Inference Architectures in Insurance Underwriting"
        ),
        size=23,
        bold=True,
        color=NAVY,
    )
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        subtitle.add_run(
            "A resource-matched comparison of direct generation, self-critique, and two-agent debate"
        ),
        size=10.5,
        color=MUTED,
    )
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(7)
    set_run(
        meta.add_run(datetime.now(UTC).strftime("%B %Y")),
        size=8.5,
        bold=True,
        color=TEAL,
    )


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    set_run(paragraph.add_run(text), size=9.5, color=INK)


def add_abstract(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT)
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), TEAL)
    border.append(left)
    p_pr.append(border)
    set_run(paragraph.add_run("Abstract. "), size=9.5, bold=True, color=NAVY)
    set_run(paragraph.add_run(text), size=9.5, color=INK)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    caption: str,
) -> None:
    caption_paragraph = document.add_paragraph()
    caption_paragraph.paragraph_format.space_before = Pt(4)
    caption_paragraph.paragraph_format.space_after = Pt(3)
    caption_paragraph.paragraph_format.keep_with_next = True
    set_run(caption_paragraph.add_run(caption), size=8, bold=True, color=NAVY)
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell._tc, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run(paragraph.add_run(header), size=8, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column, value in enumerate(values):
            if row_index % 2:
                shade(cells[column]._tc, "F7F9FA")
            paragraph = cells[column].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            set_run(paragraph.add_run(str(value)), size=8, color=INK)
    set_table_geometry(table, widths_dxa)
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(3)
    set_run(
        note.add_run("Source: experiment outputs; 95% intervals use paired case bootstrap."),
        size=7.5,
        italic=True,
        color=MUTED,
    )


def add_figure(document: Document, image_path: Path, caption: str, width=6.45) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(4)
    set_run(caption_paragraph.add_run(caption), size=8, italic=True, color=MUTED)


def percentage(value: float, digits: int = 1) -> str:
    return "n/a" if pd.isna(value) else f"{100 * value:.{digits}f}%"


def token(value: float) -> str:
    return "n/a" if pd.isna(value) else f"{value:,.0f}"


def load_results() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, dict]:
    summary_path = TABLES / "architecture_by_budget.csv"
    case_path = TABLES / "case_level_results.csv"
    difference_path = TABLES / "paired_differences.csv"
    analysis_path = OUTPUTS / "analysis_summary.json"
    required = [summary_path, case_path, difference_path, analysis_path]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise RuntimeError(f"analysis outputs are missing: {missing}")
    summary = pd.read_csv(summary_path)
    cases = pd.read_csv(case_path)
    differences = pd.read_csv(difference_path)
    analysis = json.loads(analysis_path.read_text(encoding="utf-8"))
    if analysis["successful_generations"] < analysis["expected_generations"]:
        raise RuntimeError(
            "refusing to write a results paper from an incomplete generation matrix: "
            f"{analysis['successful_generations']}/{analysis['expected_generations']}"
        )
    return summary, cases, differences, analysis


def winner_text(summary: pd.DataFrame, budget: int) -> tuple[str, float]:
    group = summary[summary["nominal_budget"] == budget]
    row = group.loc[group["accuracy"].idxmax()]
    return str(row["architecture_label"]), float(row["accuracy"])


def crossover_sentence(analysis: dict) -> str:
    clauses = []
    for row in analysis["crossovers"]:
        name = row["architecture_label"]
        value = row["crossover_budget"]
        if value is None:
            clauses.append(f"{name} did not cross the direct baseline within the tested range")
        else:
            low = row.get("crossover_ci_low")
            high = row.get("crossover_ci_high")
            interval = (
                f" (bootstrap 95% CI {low:,.0f}–{high:,.0f})"
                if low is not None and high is not None
                else ""
            )
            clauses.append(f"{name} crossed near {value:,.0f} tokens{interval}")
    return "; ".join(clauses) + "."


def build() -> Path:
    summary, case_rows, _differences, analysis = load_results()
    budgets = sorted(int(value) for value in summary["nominal_budget"].unique())
    low_budget, high_budget = budgets[0], budgets[-1]
    low_winner, low_accuracy = winner_text(summary, low_budget)
    high_winner, high_accuracy = winner_text(summary, high_budget)
    direct_low = summary[
        (summary["architecture"] == "direct") & (summary["nominal_budget"] == low_budget)
    ].iloc[0]
    direct_high = summary[
        (summary["architecture"] == "direct") & (summary["nominal_budget"] == high_budget)
    ].iloc[0]
    judge_coverage = case_rows["judge_correct"].notna().mean()
    judge_disagreement = float(analysis.get("judge_disagreement_rate", math.nan))
    support = low_winner == "Direct" and high_winner != "Direct"

    document = Document()
    configure_styles(document)
    configure_page(document)
    add_title_block(document)

    abstract = (
        f"Modern LLM systems can spend an inference budget in one long response or distribute it "
        f"across critique and debate. We tested whether the best architecture changes with budget "
        f"on {analysis['unique_cases']} deduplicated commercial-insurance underwriting cases. "
        f"A single generator model answered each case using direct generation, self-critique, or "
        f"two-agent debate under {len(budgets)} completion-token ceilings ({low_budget:,}–"
        f"{high_budget:,}). Primary correctness was task-specific and reference-exact; two blinded "
        f"LLM judges provided a secondary evidence-grounding check. {low_winner} led at the lowest "
        f"ceiling ({percentage(low_accuracy)} accuracy), while {high_winner} led at the highest "
        f"({percentage(high_accuracy)}). {crossover_sentence(analysis)} The findings "
        + (
            "support a budget-dependent routing policy rather than a universally best architecture."
            if support
            else "do not support a clean universal crossover within the tested range; architecture choice remained task- and metric-dependent."
        )
    )
    add_abstract(document, abstract)

    add_heading(document, "1. Why architecture and budget must be evaluated together")
    add_body(
        document,
        "Inference-time scaling is often discussed as if more calls necessarily buy better answers. "
        "That framing hides an allocation problem. A direct architecture can devote its full "
        "generation allowance to one coherent solution. A critique loop pays for an initial draft "
        "before it can correct it. A debate architecture pays an even larger fixed cost to create "
        "independent views and reconcile them. Under a generous budget, those extra views may expose "
        "classification or threshold errors. Under a tight budget, the same structure can fragment "
        "the answer into underspecified intermediate messages. The production question is therefore "
        "not whether debate can help in the abstract, but where its marginal benefit exceeds its "
        "coordination overhead.",
    )
    add_body(
        document,
        "Prior work shows that test-time compute should be allocated by problem difficulty [2], and "
        "recent reasoning-benchmark evidence places multi-agent methods on the compute-quality Pareto "
        "front [3]. Multi-agent debate can improve factual and reasoning performance [4], while "
        "self-refinement has reported gains across diverse tasks [5]. The evidence is not uniformly "
        "positive: a critical survey finds that prompted self-correction often fails without reliable "
        "external feedback [6]. Our contribution is a controlled domain study in multi-turn "
        "underwriting, where outputs are operational decisions rather than multiple-choice answers. "
        "We estimate architecture-by-budget curves, preserve paired cases across conditions, and "
        "report both nominal generation ceilings and measured token/latency overhead.",
    )

    add_heading(document, "2. Experimental design")
    add_heading(document, "2.1 Dataset and leakage-controlled cases", level=2)
    add_body(
        document,
        "The source is Snorkel AI's Multi-Turn Insurance Underwriting dataset [1], an Apache-2.0 "
        "collection of 380 expert-verified traces spanning appetite checks, product recommendations, "
        "policy limits, small-business eligibility, deductibles, and business classification. The "
        "380 rows are not independent tasks: they collapse to 80 company/task identifiers, typically "
        "with responses from five models. We sampled 60 unique identifiers with fixed task quotas, "
        "including every deductible and classification case. For each identifier, preprocessing "
        "selected the longest successful source trace, retained structured company fields, "
        "underwriter utterances, and tool outputs, and removed all prior assistant prose and final "
        "answers. Large SQL outputs were filtered to the case state or top lexical matches. A "
        "deterministic leakage test rejects any packet containing its reference answer.",
    )
    add_heading(document, "2.2 Architectures and resource control", level=2)
    architecture_rows = [
        ["Direct", "1", "100%", "One evidence-grounded final answer"],
        ["Self-critique", "3", "44 / 20 / 36%", "Draft → audit → corrected final"],
        [
            "Two-agent debate",
            "4",
            "25 / 25 / 20 / 30%",
            "Parallel specialists → reviewer → synthesis",
        ],
    ]
    add_table(
        document,
        ["Architecture", "Calls", "Budget allocation", "Control flow"],
        architecture_rows,
        [2100, 900, 2200, 4736],
        "Table 1. Resource-matched inference architectures.",
    )
    add_body(
        document,
        f"The generator, temperature (0.0), evidence packet, and final JSON schema were held fixed. "
        f"The independent variable was a per-case completion-token ceiling at {', '.join(f'{b:,}' for b in budgets)} "
        "tokens. Each multi-call architecture divided that ceiling across all internal calls with a "
        "64-token minimum per call; the allocations in Table 1 always sum exactly to the ceiling. "
        "This controls potential generation while allowing models to stop early. Prompt tokens, "
        "actual completion tokens, total tokens, summed call latency, finish reasons, and credential "
        "slot were recorded. Two credentials were scheduled as separate four-request pools, with "
        "parallel specialists issued concurrently. Checkpointed JSONL output made the sweep resumable, "
        "and phase deadlines capped generation plus evaluation at 7.75 hours.",
    )
    add_heading(document, "2.3 Outcomes and inference", level=2)
    add_body(
        document,
        "The primary outcome is exact operational correctness derived from the expert reference: "
        "categorical appetite/eligibility labels, numeric limits and deductibles, six-digit NAICS "
        "codes, or normalized sets of recommended lines of business. This metric is deterministic "
        "and avoids grading prose style. As a secondary check, two pointwise judges from different "
        "model families compared each answer with all accepted reference variants and scored evidence "
        "grounding from 0 to 4; a third model adjudicated correctness disagreements. Pointwise judging "
        "avoids response-order bias, although LLM-judge bias remains possible [7]. We report case-level "
        "paired bootstrap 95% intervals and fit a clustered logistic model with architecture-by-log2 "
        "budget interactions, task fixed effects, and a precomputed evidence-complexity covariate. "
        "The crossover is the first log-linear interpolation where a complex architecture's paired "
        "accuracy difference over direct changes from non-positive to positive.",
    )

    add_heading(document, "3. Results")
    add_figure(
        document,
        FIGURES / "accuracy_by_budget.png",
        "Figure 1. Exact decision accuracy by architecture and completion-token ceiling. Bands show 95% case-bootstrap intervals.",
    )
    accuracy_rows = []
    for architecture in ["direct", "self_critique", "debate"]:
        group = summary[summary["architecture"] == architecture].set_index("nominal_budget")
        accuracy_rows.append(
            [group.iloc[0]["architecture_label"]]
            + [percentage(group.loc[budget, "accuracy"]) for budget in budgets]
        )
    widths = [2400] + [(CONTENT_WIDTH_DXA - 2400) // len(budgets)] * len(budgets)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    add_table(
        document,
        ["Architecture"] + [f"{budget:,}" for budget in budgets],
        accuracy_rows,
        widths,
        "Table 2. Exact decision accuracy across nominal budgets.",
    )
    thesis_text = (
        f"The predicted dominance shift was observed: direct generation was best at {low_budget:,} "
        f"tokens ({percentage(low_accuracy)}), whereas {high_winner.lower()} was best at "
        f"{high_budget:,} ({percentage(high_accuracy)}). "
        if support
        else f"A clean dominance shift was not observed: {low_winner.lower()} led at {low_budget:,} "
        f"tokens and {high_winner.lower()} led at {high_budget:,}. "
    )
    add_body(
        document,
        thesis_text
        + crossover_sentence(analysis)
        + f" Direct accuracy changed from {percentage(direct_low['accuracy'])} to "
        f"{percentage(direct_high['accuracy'])} across the range. Confidence intervals in Figure 1 "
        "should be read as uncertainty over the sampled cases, not repeated stochastic generations; "
        "temperature was fixed to remove a second source of variance.",
    )

    resource_figure = FIGURES / "resource_tradeoff.png"
    if resource_figure.exists():
        add_figure(
            document,
            resource_figure,
            "Figure 2. Quality against measured total-token use and summed call latency. Each point is one nominal budget.",
        )
    high_rows = summary[summary["nominal_budget"] == high_budget].sort_values(
        "accuracy", ascending=False
    )
    resource_sentences = []
    for _, row in high_rows.iterrows():
        resource_sentences.append(
            f"{row['architecture_label']} used {token(row['mean_total_tokens'])} total tokens and "
            f"{row['mean_latency_seconds']:.1f} summed seconds per case at {high_budget:,}"
        )
    add_body(
        document,
        "Resource accounting changes the interpretation of nominal equality. "
        + "; ".join(resource_sentences)
        + ". Repeated inclusion of the evidence packet makes multi-call methods consume more prompt "
        "tokens even when their completion ceilings are identical. Accordingly, Figure 2 is the more "
        "relevant deployment view when gateways bill or throttle on total tokens rather than generated "
        "tokens alone.",
    )

    task_group = (
        case_rows.groupby(["task", "architecture"], as_index=False)["exact_correct"]
        .mean()
        .sort_values("exact_correct")
    )
    direct_tasks = task_group[task_group["architecture"] == "direct"]
    hardest = direct_tasks.iloc[0]["task"]
    add_body(
        document,
        f"Architecture effects were heterogeneous across task types. {hardest} had the lowest average "
        "direct accuracy in this sample, while short categorical decisions tended to leave less room "
        "for coordination gains. This pattern is consistent with difficulty-dependent test-time "
        "allocation [2], but the two business-classification cases are insufficient for a stable "
        "task-specific threshold. The adjusted logistic model and full per-task matrix are included "
        "in the repository rather than compressed into the main paper.",
    )

    add_heading(document, "4. Discussion")
    add_heading(document, "4.1 What the crossover means for system design", level=2)
    add_body(
        document,
        (
            "The results support treating architecture as a budget-conditioned routing decision. "
            if support
            else "The results caution against assuming that additional orchestration creates a predictable crossover. "
        )
        + "The direct baseline has almost no coordination overhead and remains a strong default for "
        "small ceilings or routine tasks. Self-critique is attractive when a draft can be audited "
        "against explicit rules, but intrinsic feedback can also rationalize an initial mistake [6]. "
        "Debate is most defensible when independent roles can inspect genuinely different failure "
        "modes, such as classification versus thresholds. Production routing should therefore use a "
        "joint policy over available budget, evidence complexity, and task type, with fallbacks when "
        "the gateway cannot supply the calls required by a complex plan.",
    )
    add_heading(document, "4.2 Validity and limitations", level=2)
    add_body(
        document,
        f"First, the study uses {analysis['unique_cases']} unique cases from one synthetic commercial-"
        "insurance environment; the estimated threshold is not a universal constant. Second, all "
        "architectures share one generator model and hand-designed prompts. Different base models, "
        "role prompts, or learned routers may move or remove a crossover. Third, the controlled "
        "resource is a completion-token ceiling, while actual total tokens also include repeated "
        "evidence. Cost comparisons require the internal gateway's price schedule. Fourth, a single "
        "deterministic run per cell estimates case variation but not sampling variance. Fifth, "
        f"secondary judge coverage was {percentage(judge_coverage)} and the observed disagreement "
        f"rate was {percentage(judge_disagreement)}; exact task scoring is therefore primary. Finally, "
        "selecting an evidence packet from a successful historical trace creates an answerable "
        "closed-book task but does not reproduce the original interactive tool-selection problem.",
    )

    add_heading(document, "5. Conclusion")
    add_body(
        document,
        "A token budget does not have the same value under every inference architecture. Direct "
        "generation concentrates capacity; critique and debate spend part of it on error detection "
        "and coordination. On this underwriting benchmark, the measured curves show exactly where "
        "those trade-offs pay off—or fail to—under the tested conditions. The practical deliverable "
        "is not a mandate to deploy more agents, but a reproducible method for selecting an "
        "architecture from paired quality, token, and latency evidence. Future work should replicate "
        "the sweep across generator sizes, add stochastic repeats around the estimated transition, "
        "and train a lightweight router that predicts the cheapest architecture likely to meet a "
        "target correctness level.",
    )

    add_heading(document, "References")
    references = [
        (
            "[1] Snorkel AI. Multi-Turn Insurance Underwriting (2025). ",
            "https://huggingface.co/datasets/snorkelai/Multi-Turn-Insurance-Underwriting",
        ),
        (
            "[2] C. Snell, J. Lee, K. Xu, and A. Kumar. Scaling LLM Test-Time Compute Optimally Can Be More Effective than Scaling Model Parameters (2024). ",
            "https://arxiv.org/abs/2408.03314",
        ),
        (
            "[3] F. V. Wunderlich et al. Multi-Agent Reasoning Improves Compute Efficiency: Pareto-Optimal Test-Time Scaling. ACL SRW (2026). ",
            "https://arxiv.org/abs/2605.01566",
        ),
        (
            "[4] Y. Du et al. Improving Factuality and Reasoning in Language Models through Multiagent Debate (2023). ",
            "https://arxiv.org/abs/2305.14325",
        ),
        (
            "[5] A. Madaan et al. Self-Refine: Iterative Refinement with Self-Feedback. NeurIPS 36 (2023), 46534–46594. ",
            "https://arxiv.org/abs/2303.17651",
        ),
        (
            "[6] R. Kamoi et al. When Can LLMs Actually Correct Their Own Mistakes? TACL 12 (2024), 1417–1440. ",
            "https://aclanthology.org/2024.tacl-1.78/",
        ),
        (
            "[7] J. Ye et al. Justice or Prejudice? Quantifying Biases in LLM-as-a-Judge (2024). ",
            "https://arxiv.org/abs/2410.02736",
        ),
    ]
    for label, url in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(2)
        set_run(paragraph.add_run(label), size=8, color=INK)
        add_hyperlink(paragraph, url, url)

    output = PAPER_DIR / "architecture_budget_crossover.docx"
    document.save(output)
    return output


if __name__ == "__main__":
    print(build())
